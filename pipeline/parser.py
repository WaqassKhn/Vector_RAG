import os
import io
import pandas as pd
from pathlib import Path
from typing import Dict, Any, List
from pipeline.cleaner import DocumentCleaner

# Imports with fallback handling
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class DocumentParser:
    """
    Multi-format parser capable of extracting structured text and tables from PDF, CSV, XLSX, DOCX, and TXT files.
    """

    @classmethod
    def parse_file(cls, file_path: str) -> Dict[str, Any]:
        """
        Parses a document file and returns structured text, metadata, and page breakdown.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        filename = path.name

        if ext == ".pdf":
            return cls._parse_pdf(file_path, filename)
        elif ext in [".csv", ".tsv", ".xlsx", ".xls"]:
            return cls._parse_tabular(file_path, filename, ext)
        elif ext in [".docx", ".doc"]:
            return cls._parse_docx(file_path, filename)
        elif ext in [".txt", ".md", ".log"]:
            return cls._parse_text(file_path, filename)
        else:
            # Fallback to text reading
            return cls._parse_text(file_path, filename)

    @classmethod
    def _parse_pdf(cls, file_path: str, filename: str) -> Dict[str, Any]:
        """Parse PDF documents extracting text and structured tables per page."""
        page_contents = []
        raw_text_blocks = []

        if HAS_PDFPLUMBER:
            with pdfplumber.open(file_path) as pdf:
                for page_idx, page in enumerate(pdf.pages):
                    page_num = page_idx + 1
                    page_text = page.extract_text() or ""
                    
                    # Extract tables from page
                    tables = page.extract_tables()
                    md_tables = []
                    for t in tables:
                        md_t = DocumentCleaner.format_raw_table_to_markdown(t)
                        if md_t:
                            md_tables.append(md_t)
                    
                    combined = page_text
                    if md_tables:
                        combined += "\n\n### Extracted Tables:\n" + "\n\n".join(md_tables)
                    
                    cleaned = DocumentCleaner.clean_document_text(combined)
                    page_contents.append({
                        "page_number": page_num,
                        "raw_text": page_text,
                        "cleaned_text": cleaned
                    })
                    raw_text_blocks.append(cleaned)
        elif HAS_PYPDF:
            reader = pypdf.PdfReader(file_path)
            for page_idx, page in enumerate(reader.pages):
                page_num = page_idx + 1
                page_text = page.extract_text() or ""
                cleaned = DocumentCleaner.clean_document_text(page_text)
                page_contents.append({
                    "page_number": page_num,
                    "raw_text": page_text,
                    "cleaned_text": cleaned
                })
                raw_text_blocks.append(cleaned)
        else:
            raise ImportError("Neither pdfplumber nor pypdf is installed for PDF parsing.")

        full_cleaned_text = "\n\n--- Page Break ---\n\n".join(raw_text_blocks)
        
        return {
            "filename": filename,
            "file_type": "pdf",
            "pages": page_contents,
            "full_text": full_cleaned_text,
            "num_pages": len(page_contents)
        }

    @classmethod
    def _parse_tabular(cls, file_path: str, filename: str, ext: str) -> Dict[str, Any]:
        """Parse CSV or Excel spreadsheets into formatted Markdown tables."""
        if ext == ".csv":
            df = pd.read_csv(file_path)
        elif ext == ".tsv":
            df = pd.read_csv(file_path, sep="\t")
        else:
            df = pd.read_excel(file_path)

        # Convert DataFrame to records matrix including header
        table_matrix = [list(df.columns)] + df.values.tolist()
        md_table = DocumentCleaner.format_raw_table_to_markdown(table_matrix)
        
        cleaned_text = f"# Sheet: {filename}\n\n" + md_table
        
        return {
            "filename": filename,
            "file_type": ext.lstrip("."),
            "pages": [{"page_number": 1, "raw_text": str(table_matrix), "cleaned_text": cleaned_text}],
            "full_text": cleaned_text,
            "num_pages": 1
        }

    @classmethod
    def _parse_docx(cls, file_path: str, filename: str) -> Dict[str, Any]:
        """Parse Microsoft Word documents."""
        if not HAS_DOCX:
            raise ImportError("python-docx is required to parse .docx files.")

        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Extract tables if any
        table_mds = []
        for table in doc.tables:
            table_matrix = []
            for row in table.rows:
                table_matrix.append([cell.text for cell in row.cells])
            md_t = DocumentCleaner.format_raw_table_to_markdown(table_matrix)
            if md_t:
                table_mds.append(md_t)

        combined_text = "\n\n".join(paragraphs)
        if table_mds:
            combined_text += "\n\n### Document Tables:\n" + "\n\n".join(table_mds)

        cleaned_text = DocumentCleaner.clean_document_text(combined_text)

        return {
            "filename": filename,
            "file_type": "docx",
            "pages": [{"page_number": 1, "raw_text": combined_text, "cleaned_text": cleaned_text}],
            "full_text": cleaned_text,
            "num_pages": 1
        }

    @classmethod
    def _parse_text(cls, file_path: str, filename: str) -> Dict[str, Any]:
        """Parse plain text / markdown files."""
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            raw_content = f.read()

        cleaned_text = DocumentCleaner.clean_document_text(raw_content)

        return {
            "filename": filename,
            "file_type": "txt",
            "pages": [{"page_number": 1, "raw_text": raw_content, "cleaned_text": cleaned_text}],
            "full_text": cleaned_text,
            "num_pages": 1
        }
